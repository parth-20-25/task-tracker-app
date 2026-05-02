from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\Dell\Downloads\TaskTrackerApp\deliverables\PARC_SOP_Standard_Operating_Procedure.docx")

NAVY = "173753"
STEEL = "4F6D7A"
GOLD = "B8860B"
LIGHT_BLUE = "EAF1F6"
LIGHT_STEEL = "EDF2F7"
LIGHT_GOLD = "F7F2E7"
LIGHT_RED = "FAEAEA"
DARK = RGBColor(28, 37, 44)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_background(paragraph, fill):
    p_pr = paragraph._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, *, font="Calibri", size=10.5, bold=False, italic=False, color=None, all_caps=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.all_caps = all_caps
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, *, before=0, after=6, line=1.15, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    format_run(run, **kwargs)
    return run


def add_body_paragraph(document, text):
    p = document.add_paragraph()
    style_paragraph(p, before=0, after=7, line=1.2)
    add_text(p, text, font="Calibri", size=10.5, color="2C3E50")
    return p


def add_bullet(document, text):
    p = document.add_paragraph()
    style_paragraph(p, before=0, after=4, line=1.15)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    add_text(p, "- ", font="Calibri", size=10.25, bold=True, color=NAVY)
    add_text(p, text, font="Calibri", size=10.25, color="31424F")
    return p


def add_numbered(document, text):
    p = document.add_paragraph()
    style_paragraph(p, before=0, after=4, line=1.15)
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_text(p, text, font="Calibri", size=10.25, color="31424F")
    return p


def add_heading_band(document, title, subtitle=None):
    p = document.add_paragraph()
    style_paragraph(p, before=14, after=4)
    set_paragraph_background(p, LIGHT_BLUE)
    add_text(p, f"  {title}", font="Cambria", size=15, bold=True, color=NAVY)
    if subtitle:
        s = document.add_paragraph()
        style_paragraph(s, before=0, after=10)
        add_text(s, subtitle, font="Calibri", size=10.5, italic=True, color=STEEL)
    return p


def add_subheading(document, title):
    p = document.add_paragraph()
    style_paragraph(p, before=10, after=4)
    add_text(p, title, font="Cambria", size=12.5, bold=True, color=NAVY)
    return p


def add_kv_table(document, rows, widths=(1.9, 4.9), shade_first=True):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        if shade_first:
            set_cell_shading(cells[0], LIGHT_STEEL)
        p0 = cells[0].paragraphs[0]
        style_paragraph(p0, before=0, after=0, line=1.0)
        add_text(p0, key, font="Calibri", size=10, bold=True, color=NAVY)
        p1 = cells[1].paragraphs[0]
        style_paragraph(p1, before=0, after=0, line=1.1)
        add_text(p1, value, font="Calibri", size=10, color="2C3E50")
    document.add_paragraph()
    return table


def add_matrix_table(document, headers, rows, widths=None, fill_header=NAVY):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, cell in enumerate(hdr.cells):
        set_cell_shading(cell, fill_header)
        set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(p, headers[idx], font="Calibri", size=9.5, bold=True, color="FFFFFF")
        if widths:
            cell.width = Inches(widths[idx])
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            if widths:
                cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, before=0, after=0, line=1.08, alignment=align)
            add_text(p, str(value), font="Calibri", size=9.3, color="2C3E50")
    document.add_paragraph()
    return table


def add_callout(document, title, body, fill=LIGHT_GOLD):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p1 = cell.paragraphs[0]
    style_paragraph(p1, before=0, after=3, line=1.0)
    add_text(p1, title, font="Cambria", size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    style_paragraph(p2, before=0, after=0, line=1.1)
    add_text(p2, body, font="Calibri", size=10, color="2C3E50")
    document.add_paragraph()
    return table


def add_process_flow(document, title, steps):
    add_subheading(document, title)
    table = document.add_table(rows=2, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for idx, step in enumerate(steps):
        top = table.rows[0].cells[idx]
        bottom = table.rows[1].cells[idx]
        set_cell_shading(top, NAVY if idx % 2 == 0 else STEEL)
        set_cell_shading(bottom, LIGHT_STEEL)
        for cell in (top, bottom):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, start=90, bottom=100, end=90)
        ptop = top.paragraphs[0]
        style_paragraph(ptop, before=0, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(ptop, f"Step {idx + 1}", font="Calibri", size=9.5, bold=True, color="FFFFFF")
        pbot = bottom.paragraphs[0]
        style_paragraph(pbot, before=0, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(pbot, step, font="Calibri", size=9.2, bold=True, color=NAVY)
    document.add_paragraph()
    return table


def add_page_break(document):
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = "Cambria"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Cambria")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(footer_p, before=0, after=0, line=1.0)
    add_text(
        footer_p,
        "PARC | Project Assignment, Reporting & Control System | Controlled Internal Operating Document",
        font="Calibri",
        size=8.5,
        color=STEEL,
    )


def build_cover(document):
    top_band = document.add_table(rows=1, cols=1)
    top_band.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_band.style = "Table Grid"
    cell = top_band.rows[0].cells[0]
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=180, start=180, bottom=180, end=180)
    p = cell.paragraphs[0]
    style_paragraph(p, before=0, after=2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p, "STANDARD OPERATING PROCEDURE", font="Cambria", size=20, bold=True, color="FFFFFF")
    p2 = cell.add_paragraph()
    style_paragraph(p2, before=0, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(p2, "Industrial Workflow Governance and Platform Usage Manual", font="Calibri", size=11, color="FFFFFF")

    hero = document.add_paragraph()
    style_paragraph(hero, before=26, after=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(hero, "PARC", font="Cambria", size=30, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    style_paragraph(subtitle, before=0, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(subtitle, "Project Assignment, Reporting & Control System", font="Calibri", size=15, bold=True, color=STEEL)

    strap = document.add_paragraph()
    style_paragraph(strap, before=6, after=12, line=1.3, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        strap,
        "Executive-grade operating standard for controlled task execution, departmental coordination, approval discipline, escalation governance, and performance visibility across industrial operations.",
        font="Calibri",
        size=11.5,
        color="2C3E50",
    )

    add_callout(
        document,
        "Document Positioning",
        "This SOP defines how PARC shall be used as a mission-critical control platform across design, planning, production, quality, maintenance, stores, administration, and future operational functions. It is intended for founders, directors, department heads, line managers, supervisors, engineers, designers, and shop-floor users.",
        fill=LIGHT_BLUE,
    )

    add_kv_table(
        document,
        [
            ("Document Code", "PARC-SOP-001"),
            ("Version", "1.0"),
            ("Classification", "Controlled Internal Document"),
            ("Intended Use", "Operations governance, onboarding, implementation control, audit readiness"),
            ("Primary Audience", "Leadership, department management, supervisors, engineers, designers, operators, end users"),
            ("Review Cycle", "Quarterly or upon major process change"),
        ],
    )

    note = document.add_paragraph()
    style_paragraph(note, before=18, after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(note, "Prepared for enterprise operational deployment and management review.", font="Calibri", size=9.5, italic=True, color=STEEL)


def build_contents(document):
    add_page_break(document)
    add_heading_band(document, "Document Control", "Governance, ownership, applicability, and control requirements.")
    add_kv_table(
        document,
        [
            ("Document Title", "Standard Operating Procedure for PARC (Project Assignment, Reporting & Control System)"),
            ("Process Owner", "Operations Excellence / Digital Process Governance"),
            ("Approval Authority", "Founders, Directors, Plant Head, Department Heads"),
            ("System Owner", "PARC Platform Administration with Department Functional Owners"),
            ("Applies To", "All active users across configured departments and any future department onboarded to PARC"),
            ("Supersedes", "Informal task follow-up via verbal, spreadsheet, or disconnected messaging methods"),
            ("Source of Truth Rule", "Where operational status exists in PARC, PARC shall be treated as the auditable system of record"),
            ("Deviation Rule", "Any workflow deviation must be recorded inside PARC through task remarks, issue reporting, rework, or escalation records"),
        ],
    )

    add_subheading(document, "Section Index")
    sections = [
        "1. Purpose",
        "2. Scope",
        "3. System Overview",
        "4. Business Need",
        "5. Roles and Responsibilities",
        "6. User Hierarchy Structure",
        "7. Workflow Process Flow",
        "8. Module-wise Explanation",
        "9. Department-wise Usage",
        "10. Assignment Management Process",
        "11. Rework Handling Process",
        "12. Escalation Management",
        "13. Analytics and Performance Monitoring",
        "14. Approval and Reporting Structure",
        "15. Rules for Data Entry",
        "16. User and Management Responsibilities",
        "17. Best Practices and Compliance",
        "18. Common Mistakes to Avoid",
        "19. Escalation Matrix",
        "20. Operational Benefits and Business Outcomes",
        "21. Future Scalability",
        "22. Conclusion",
    ]
    for item in sections:
        add_bullet(document, item)


def populate_sections(document):
    add_heading_band(document, "Purpose")
    add_body_paragraph(
        document,
        "The purpose of PARC is to establish one disciplined operating system for assigning work, tracking execution, validating completion, controlling rework, escalating delays, and converting daily transactional activity into reliable management intelligence. This SOP defines the expected operating behavior around that system so that work is not only performed, but controlled, attributable, reviewable, and decision-ready.",
    )
    add_body_paragraph(
        document,
        "PARC exists to eliminate ambiguity in ownership, reduce hidden delays between departments, prevent unmanaged task duplication, formalize approvals, and ensure that every critical work item moves through a visible workflow rather than through informal follow-up. In practical terms, PARC transforms operational activity from person-dependent coordination into process-dependent control.",
    )

    add_heading_band(document, "Scope")
    add_body_paragraph(
        document,
        "This SOP applies to all work executed inside PARC, including structured department workflow tasks, controlled project assignments, fixture-stage execution, verification activity, rework handling, issue reporting, escalation action, batch data uploads, operational reports, and management analytics. It is valid for current departments such as Design, PPC, Production, Quality, Maintenance, Stores, and Admin, and it shall also govern future departments added to the platform.",
    )
    add_body_paragraph(
        document,
        "The SOP covers both operational usage and management usage. It therefore applies equally to the employee who executes the work, the supervisor who verifies it, the department head who monitors capacity and compliance, and the director or founder who relies on the resulting analytics to assess control, responsiveness, and operational stability.",
    )

    add_heading_band(document, "System Overview")
    add_body_paragraph(
        document,
        "PARC stands for Project Assignment, Reporting & Control System. It is an industrial workflow management platform designed to handle assignment ownership, departmental execution, approval flow, rework discipline, escalation routing, auditability, and executive visibility from one shared operating environment.",
    )
    add_matrix_table(
        document,
        ["Control Area", "What PARC Does", "Why It Matters"],
        [
            ("Assignment Control", "Creates accountable ownership for every work item.", "No task should exist without a named owner, origin, deadline, and status."),
            ("Workflow Control", "Moves work through stage-based execution and approval.", "Prevents process jumping, manual ambiguity, and uncontrolled handoffs."),
            ("Verification Control", "Routes completed work into supervisor or quality review.", "Ensures closure is based on acceptance, not only on self-declared completion."),
            ("Rework Control", "Captures rejection, remarks, and return-to-work cycles.", "Makes quality loss visible and measurable instead of hidden."),
            ("Escalation Control", "Raises unattended work according to priority and timing rules.", "Protects delivery commitments before delays become failures."),
            ("Performance Control", "Generates role-appropriate analytics and reports.", "Allows leaders to act on facts rather than anecdotal updates."),
        ],
        widths=[1.55, 2.7, 2.35],
    )
    add_callout(
        document,
        "Operating Principle",
        "PARC is not a passive tracker. It is a control mechanism. A task only has management value when the owner, status, proof, approval, and timing signals are all captured correctly inside the system.",
        fill=LIGHT_GOLD,
    )

    add_heading_band(document, "Business Need")
    add_body_paragraph(
        document,
        "Industrial operations break down when coordination depends on memory, verbal follow-up, personal spreadsheets, or disconnected chat updates. In such environments, the same work can be assigned twice, approvals can be assumed instead of recorded, rework can be repeated without cause analysis, and leadership receives status after the risk has already materialized. PARC is the control response to those failure modes.",
    )
    add_body_paragraph(
        document,
        "The platform addresses a specific business need: create a single system where execution, verification, delay, rejection, escalation, and performance evidence are linked together. This makes it possible to answer high-value management questions at any moment: What is pending? Who owns it? Which department is constraining flow? Where is rework recurring? Which stage is unstable? Which teams are reliable against deadlines? Which risks require intervention today?",
    )
    add_callout(
        document,
        "Examples of Problems PARC Solves",
        "A design stage approved by a supervisor should automatically release the next stage instead of waiting for another manual reminder. A production task should not be re-assigned to a second person while the first active assignment still exists. A quality rejection should create a visible rework cycle with remarks and timing impact, not an undocumented verbal instruction.",
        fill=LIGHT_BLUE,
    )

    add_heading_band(document, "Roles and Responsibilities")
    add_matrix_table(
        document,
        ["Role Group", "Primary Use in PARC", "Expected Responsibility"],
        [
            ("Founders / Directors", "Review enterprise visibility, reports, and strategic analytics.", "Use PARC to validate control maturity, delivery reliability, and departmental accountability."),
            ("Plant Head / Department Heads", "Monitor cross-team performance and approve governance decisions.", "Ensure departmental adoption, corrective action, and compliance with reporting discipline."),
            ("Line Managers", "Create and assign structured workflow tasks, manage queues, review exceptions.", "Maintain balanced workload, enforce deadlines, and keep the system operationally current."),
            ("Shift Incharges / Supervisors", "Assign work, review submitted tasks, return rework, close execution loops.", "Verify actual completion against standard, not assumption."),
            ("Quality Inspectors", "Approve quality-sensitive tasks and review rejections.", "Protect process integrity by documenting acceptance and nonconformity inside the platform."),
            ("Engineers / Designers / Operators", "Execute assigned tasks, upload proof, update status, respond to rework.", "Keep task data complete, timely, and truthful."),
            ("Admin / System Coordinators", "Maintain users, departments, permissions, workflows, escalation rules.", "Protect master data integrity and platform governance."),
        ],
        widths=[1.55, 2.6, 2.45],
    )

    add_heading_band(document, "User Hierarchy Structure")
    add_body_paragraph(
        document,
        "PARC follows a hierarchy-based permission model. Users do not see everything by default; they see what is relevant to their role, department, and reporting scope. This ensures visibility is intentional, not accidental, while still preserving upward transparency for management and executive review.",
    )
    add_matrix_table(
        document,
        ["Seeded Role", "Level", "Scope", "Operational Meaning"],
        [
            ("Admin", "1", "Global", "Full platform control across users, departments, workflows, analytics, and reports."),
            ("Plant Head", "2", "Department / enterprise leadership", "High-level control, broader task visibility, escalation and analytics oversight."),
            ("Line Manager", "3", "Team", "Department execution leadership, assignment control, review and reporting."),
            ("Shift Incharge", "4", "Team", "Daily work allocation and first-line verification control."),
            ("Quality Inspector", "4", "Department", "Quality-stage review, rework disposition, reporting support."),
            ("Maintenance Engineer", "5", "Self", "Execution ownership with proof upload and self-performance visibility."),
            ("Operator", "6", "Self", "Task execution, status update, proof submission, and response to rework."),
        ],
        widths=[1.7, 0.7, 1.1, 3.0],
    )
    add_body_paragraph(
        document,
        "The hierarchy is supported by role-based permissions such as task assignment, verification, quality approval, workflow administration, escalation rule management, report access, and tiered analytics access. The result is a controlled environment where each person acts within authority while management can still review the entire operational picture.",
    )

    add_heading_band(document, "Workflow Process Flow")
    add_process_flow(
        document,
        "Core Operational Flow",
        [
            "Work Identified",
            "Task Assigned",
            "Execution Started",
            "Proof Submitted",
            "Review / Approval",
            "Closed or Rework",
        ],
    )
    add_body_paragraph(
        document,
        "PARC uses a controlled lifecycle for operational work: a task is created, assigned to a named owner, executed against a required due date, submitted with supporting proof where applicable, reviewed by a supervisor or quality authority, and then either closed or returned for rework. No step should be skipped, and a later step must not be treated as complete when the prior control point has not been recorded.",
    )
    add_body_paragraph(
        document,
        "For structured workflow tasks, PARC distinguishes between operational work that belongs inside department analytics and custom intervention tasks that are intentionally kept visible but excluded from workflow-performance scoring. This distinction is important because it prevents exceptional management interventions from distorting standard operational analytics.",
    )
    add_process_flow(
        document,
        "Representative Design Fixture Flow",
        [
            "Concept",
            "DAP",
            "3D Finish",
            "2D Finish",
            "Release",
        ],
    )
    add_body_paragraph(
        document,
        "The current seeded design workflow illustrates PARC's stage-based control logic. Once a stage-linked task is approved, the workflow can advance to the next configured stage automatically, which removes dependency on manual follow-up and ensures that downstream work becomes available only after the preceding stage is genuinely approved.",
    )

    add_heading_band(document, "Module-wise Explanation")
    add_matrix_table(
        document,
        ["Module", "Operational Purpose", "Typical Users"],
        [
            ("Dashboard", "Displays workload, pending reviews, and control metrics.", "All users, with management summary views for leaders."),
            ("My Tasks", "Personal execution queue with status updates and proof submission.", "Designers, engineers, operators, shop-floor employees."),
            ("Team Tasks (Kanban)", "Department-level board for monitoring active operational flow.", "Managers, supervisors, team leaders."),
            ("Verification Queue", "Controlled review area for submitted work awaiting approval.", "Supervisors, line managers, quality inspectors."),
            ("Notifications", "Time-sensitive communication for assignment, review, escalation, and action prompts.", "All users."),
            ("Issues", "Create, assign, and track controlled issue reports.", "Any department raising or resolving exceptions."),
            ("Batches", "Review uploaded fixture batches and manage data intake safely.", "Design data owners, department coordinators."),
            ("Reports", "Generate structured department reports and active progress views.", "Department heads, managers, auditors, directors."),
            ("Analytics", "Backend-verified KPIs including rework, deadline, efficiency, workflow health, and predictive risk.", "Leadership and authorized managers."),
            ("Admin Panel", "Manage departments, users, roles, workflows, escalation rules, shifts, and machines.", "System administrators and authorized leadership."),
        ],
        widths=[1.35, 3.2, 2.0],
    )

    add_heading_band(document, "Department-wise Usage")
    add_matrix_table(
        document,
        ["Department", "How PARC Is Used", "Control Benefit"],
        [
            ("Design", "Project scope intake, fixture registration, staged execution from Concept to Release, proof and approval tracking.", "Controls engineering handoffs and makes design rework measurable."),
            ("PPC", "Planning-driven assignments, follow-up on due dates, cross-department coordination, readiness tracking.", "Aligns planning commitments with live execution status."),
            ("Production", "Shift validation, batch verification, machine readiness, execution follow-up.", "Improves shop-floor accountability and reduces silent backlog accumulation."),
            ("Quality", "Inspection, rejection review, final audit, quality-gate approvals.", "Prevents nonconforming work from being treated as complete."),
            ("Maintenance", "Preventive checks, breakdown verification, machine recovery tasks.", "Shortens response cycles and exposes repeat equipment instability."),
            ("Stores", "Material or tooling support tasks, issue response, handoff confirmation.", "Makes supply-side delays visible to the operating chain."),
            ("Admin", "User governance, master data upkeep, workflow and escalation administration.", "Protects platform reliability and audit readiness."),
        ],
        widths=[1.0, 3.65, 1.9],
    )
    add_callout(
        document,
        "Cross-Department Example",
        "A new fixture may start in Design, require PPC planning alignment, proceed into Production execution, trigger Quality inspection, and finally depend on Maintenance or Stores support when a machine or material issue appears. PARC preserves this chain through linked ownership and visible status rather than disconnected departmental updates.",
        fill=LIGHT_GOLD,
    )

    add_heading_band(document, "Assignment Management Process")
    add_numbered(document, "1. The assigner shall select the correct department, task type, template, priority, assignee, and deadline before issuing any live task.")
    add_numbered(document, "2. Template-driven department workflow tasks shall be used for repeatable operational work that must appear in analytics and management reporting.")
    add_numbered(document, "3. Custom tasks shall be used only for controlled intervention work, exceptions, or management actions that should remain visible without distorting workflow KPI baselines.")
    add_numbered(document, "4. Proof and approval requirements must be set deliberately. Critical work should not be configured without the appropriate verification gate.")
    add_numbered(document, "5. An assignment is not considered operationally released until it appears under the named assignee and the owner has a clear due date and execution brief.")
    add_body_paragraph(
        document,
        "PARC also prevents duplicate active assignment in controlled workflow scenarios. For example, if a fixture stage already has an active task, the same stage should not be assigned again until the active task is closed or cancelled. This is a critical control because duplicate ownership is one of the most common causes of waste, confusion, and accountability failure in manufacturing administration.",
    )
    add_matrix_table(
        document,
        ["Task Status", "Meaning", "Expected User Behavior"],
        [
            ("created", "Prepared but not operationally released.", "Do not treat as executable until assignment is completed."),
            ("assigned", "Owned and awaiting start.", "Assignee must acknowledge and start within expected response time."),
            ("in_progress", "Execution has started.", "User must work against the defined requirement and due date."),
            ("on_hold", "Temporarily paused.", "Pause reason should be meaningful and visible to management."),
            ("under_review", "Submitted for verification.", "No self-closure; task awaits review action."),
            ("rework", "Returned for correction.", "Owner must address remarks and re-submit."),
            ("closed", "Approved and operationally complete.", "Downstream handoff or analytics recognition may proceed."),
            ("cancelled", "Withdrawn by authority.", "Should be used sparingly and never to hide poor control."),
        ],
        widths=[1.15, 1.95, 3.15],
    )

    add_heading_band(document, "Rework Handling Process")
    add_body_paragraph(
        document,
        "Rework in PARC is not a side conversation. It is a formal operating event. When a supervisor or quality reviewer rejects work, the task returns to the responsible owner with remarks, the rejection count increases, and the rework cycle becomes visible for both operational follow-up and analytics. This is essential because hidden rework destroys planning reliability and masks the true cost of poor execution.",
    )
    add_numbered(document, "1. Reviewer records rejection inside PARC with clear remarks linked to the task or stage.")
    add_numbered(document, "2. Task status moves to rework and returns to the responsible user.")
    add_numbered(document, "3. Owner corrects the work, updates proof if applicable, and re-submits for review.")
    add_numbered(document, "4. Reviewer evaluates the corrected work and either approves closure or raises another rework event.")
    add_body_paragraph(
        document,
        "A practical example is a design fixture released with an incomplete drawing condition. Instead of informal verbal correction, the reviewer returns the task for rework with specific remarks. The rework becomes countable, attributable to a stage, time-linked to the delay, and visible later in management analytics. This protects both accountability and learning.",
    )

    add_heading_band(document, "Escalation Management")
    add_body_paragraph(
        document,
        "Escalation in PARC is designed to protect delivery, not to punish users. Escalation rules allow the organization to define how long a task can remain unattended at a given priority before a higher authority is notified or expected to intervene. The system supports escalation levels and next-escalation timing so that sensitive work does not disappear into passive queues.",
    )
    add_body_paragraph(
        document,
        "Escalation should be used for missed response, stalled execution, missed verification, repeated rework, and cross-department dependency blockage. It should not be used to compensate for poor task creation, vague instructions, or careless data entry. Management must therefore treat escalation as a control outcome, not as a substitute for disciplined planning.",
    )
    add_callout(
        document,
        "Representative Escalation Example",
        "A critical machine-readiness task issued to Maintenance for an urgent production run should not wait until the next review meeting for action. If the task remains unresolved beyond the defined priority threshold, PARC escalation must bring it to the supervisor, manager, or department head in time for operational recovery.",
        fill=LIGHT_RED,
    )

    add_heading_band(document, "Analytics and Performance Monitoring")
    add_body_paragraph(
        document,
        "PARC converts daily task behavior into operational intelligence. Authorized users can review analytics related to self-performance, department performance, cross-department comparisons, deadline honesty, rework exposure, efficiency, workflow health, and predictive risk signals. Because the analytics are tied to controlled workflow events rather than informal updates, they have management credibility.",
    )
    add_matrix_table(
        document,
        ["Analytics Area", "What Leadership Can See", "Why It Matters"],
        [
            ("Overview", "Total tasks, approval rate, overdue rate, rework rate, department comparisons.", "Provides a fast operational pulse."),
            ("User Performance", "Approved tasks, on-time closures, overdue count, rework count, score and ranking.", "Supports coaching, recognition, and intervention."),
            ("Deadline Honesty", "Difference between promised and actual completion behavior.", "Exposes unrealistic planning or weak commitment discipline."),
            ("Workflow Health", "Composite view of quality, timeliness, stability, and control health.", "Highlights where process maturity is weakening."),
            ("Predictive Insights", "Rework probability, delay risk, likely bottleneck stage, projected remaining effort.", "Moves management from reactive review to preventive action."),
            ("Reports", "Department reports, active scope progress, export-ready summaries.", "Supports review meetings, audits, and decision packs."),
        ],
        widths=[1.45, 3.2, 1.85],
    )
    add_body_paragraph(
        document,
        "Management trust in PARC depends on disciplined usage. If deadlines are not maintained, proof is not uploaded, or statuses are not updated on time, analytics will still function technically but will lose managerial value. The platform is only as trustworthy as the operating behavior around it. This SOP therefore treats data discipline as a governance requirement, not as a clerical preference.",
    )

    add_heading_band(document, "Approval and Reporting Structure")
    add_body_paragraph(
        document,
        "PARC separates execution from acceptance. The person who completes work is not the same control point as the person who verifies it. Submitted tasks enter a verification queue where authorized reviewers can approve or reject them. Quality-sensitive tasks can require additional quality approval, ensuring that closure is based on standard compliance, not only on completion claims.",
    )
    add_numbered(document, "1. Assignee completes the task and uploads proof if required.")
    add_numbered(document, "2. Task moves to under review and appears in the verification queue.")
    add_numbered(document, "3. Supervisor or line manager verifies suitability, completeness, and remarks.")
    add_numbered(document, "4. If quality approval is required, the task remains subject to the quality gate before closure.")
    add_numbered(document, "5. Only after approval does the task move to closed and become valid for final reporting and workflow progression.")
    add_body_paragraph(
        document,
        "This control structure is especially important in industrial environments where a missed inspection point or premature closure can create downstream defects, planning inaccuracies, and customer-facing risk. PARC therefore uses approval as an operating control, not a ceremonial step.",
    )

    add_heading_band(document, "Rules for Data Entry")
    add_matrix_table(
        document,
        ["Rule", "Required Standard", "Reason"],
        [
            ("Owner fields", "Every task must have a valid assignee and assigning authority.", "Prevents anonymous work and blame diffusion."),
            ("Department selection", "Task must be issued under the correct department context.", "Keeps visibility, workflow, and analytics accurate."),
            ("Deadline quality", "Use realistic, reviewable due dates rather than placeholder dates.", "Protects deadline analytics from artificial distortion."),
            ("Description clarity", "Brief must define what to do, what standard applies, and what completion means.", "Reduces avoidable rework and clarification loss."),
            ("Remarks discipline", "Use remarks for meaningful exceptions, rejection reasons, and hold causes.", "Creates an audit trail that management can act on."),
            ("Proof discipline", "Upload proof where required before seeking approval.", "Supports evidence-based review."),
            ("Status accuracy", "Update status at the actual point of change, not after memory-based delay.", "Preserves real-time visibility."),
        ],
        widths=[1.35, 3.15, 2.0],
    )
    add_callout(
        document,
        "Data Entry Principle",
        "If an operational fact is important enough to be discussed, it is important enough to be recorded in PARC. The system cannot protect the business from information that never enters the system.",
        fill=LIGHT_BLUE,
    )

    add_heading_band(document, "User Responsibilities")
    add_bullet(document, "Accept only those assignments that belong to the user and raise mismatches immediately.")
    add_bullet(document, "Begin work promptly and keep status current.")
    add_bullet(document, "Submit complete proof where proof is mandatory.")
    add_bullet(document, "Read rework remarks carefully and correct the stated issue before re-submission.")
    add_bullet(document, "Use issue reporting when execution is blocked by another department, missing material, machine condition, or unclear input.")
    add_bullet(document, "Treat PARC updates as part of the job, not as separate administrative work.")

    add_heading_band(document, "Management Responsibilities")
    add_bullet(document, "Ensure every team member is trained on correct task handling, review flow, and evidence expectations.")
    add_bullet(document, "Use department workflow templates for standard work so analytics remain meaningful.")
    add_bullet(document, "Review overdue items, verification queues, rework patterns, and escalations as part of daily control.")
    add_bullet(document, "Do not bypass PARC through side-channel approvals or verbal closures.")
    add_bullet(document, "Correct poor data quality immediately; do not allow false closure or stale tasks to remain in the system.")
    add_bullet(document, "Use analytics for action, not only for reporting. Visible risk without follow-up is still uncontrolled risk.")

    add_heading_band(document, "Best Practices")
    add_bullet(document, "Use workflow templates for recurring operational work and reserve custom tasks for exceptions or interventions.")
    add_bullet(document, "Issue tasks as close as possible to the point of execution need, with clear deadlines and defined success criteria.")
    add_bullet(document, "Keep review remarks factual, specific, and correction-oriented.")
    add_bullet(document, "Use issue reporting early when a dependency threatens completion.")
    add_bullet(document, "Close the loop on every escalation by recording what action was taken.")
    add_bullet(document, "Review rework trends by department and stage rather than only by individual occurrence.")

    add_heading_band(document, "Compliance Rules")
    add_bullet(document, "No critical operational work shall be treated as complete unless closed through the PARC review flow.")
    add_bullet(document, "No supervisor shall approve work without sufficient proof or review basis where proof is required.")
    add_bullet(document, "No department shall maintain a parallel private tracker as its primary operational source of truth when PARC is active for that work.")
    add_bullet(document, "Role-based access and hierarchical visibility shall not be bypassed by credential sharing or informal approvals.")
    add_bullet(document, "Escalation, rejection, and issue records shall not be suppressed for the sake of presentation quality.")

    add_heading_band(document, "Common Mistakes to Avoid")
    add_matrix_table(
        document,
        ["Common Mistake", "Why It Is Harmful", "Correct Practice"],
        [
            ("Closing by assumption", "Makes reporting unreliable and hides incomplete work.", "Close only after formal review."),
            ("Wrong department tagging", "Distorts visibility, reporting, and escalation ownership.", "Issue under the correct departmental context."),
            ("Using vague descriptions", "Causes rework and repeated clarification.", "Define expected output and acceptance condition."),
            ("Delaying status updates", "Breaks real-time control and weakens escalation logic.", "Update when the status actually changes."),
            ("Creating duplicate assignments", "Causes ownership conflict and wasted effort.", "Check live status and workflow control before re-assigning."),
            ("Treating rework as informal correction", "Hides quality loss and denies learning.", "Use the rework cycle inside PARC."),
        ],
        widths=[1.75, 2.65, 1.85],
    )

    add_heading_band(document, "Escalation Matrix")
    add_matrix_table(
        document,
        ["Scenario", "Trigger", "Primary Owner", "Escalated To", "Expected Response"],
        [
            ("Task not started", "Assigned task remains unattended beyond allowed response window.", "Assignee", "Supervisor / Shift Incharge", "Clarify blockage or force action."),
            ("Task overdue", "Due date missed with task still open.", "Assignee and immediate manager", "Line Manager / Department Head", "Recover schedule and reset priorities."),
            ("Verification delay", "Completed work remains pending review for too long.", "Reviewer", "Department Head / Quality authority", "Clear queue and protect downstream flow."),
            ("Repeated rework", "Same task or stage returns with repeated rejection events.", "Assignee and reviewer", "Line Manager / Functional Head", "Correct root cause, not only symptom."),
            ("Cross-department dependency block", "Material, machine, drawing, or planning dependency stops work.", "Owning department and dependent department", "Concerned managers", "Intervene jointly and document decision."),
            ("Critical operational risk", "High-priority work threatens delivery, quality, or customer commitment.", "Responsible department lead", "Plant Head / Director", "Immediate decision and visible follow-through."),
        ],
        widths=[1.6, 1.75, 1.3, 1.35, 1.9],
    )

    add_heading_band(document, "Operational Benefits")
    add_body_paragraph(
        document,
        "When used correctly, PARC delivers operational discipline at multiple levels simultaneously. On the shop floor, it clarifies who is doing what and by when. At supervisory level, it provides clean review queues and visible backlogs. At departmental level, it exposes rework, delay, and handoff performance. At executive level, it converts operational noise into structured decision support.",
    )
    add_matrix_table(
        document,
        ["Benefit Category", "Direct Operational Effect"],
        [
            ("Visibility", "Live task status, workload, and pending approvals remain visible to the right authority."),
            ("Accountability", "Every activity is attributable to a user, department, and time event."),
            ("Quality Control", "Rework and approval gating prevent false closure."),
            ("Speed of Decision", "Escalation and analytics reduce dependency on manual status collection."),
            ("Audit Readiness", "System records create a reviewable chain of actions and decisions."),
            ("Scalability", "Standardized workflows can be reused across departments as operations expand."),
        ],
        widths=[2.1, 4.5],
    )

    add_heading_band(document, "Expected Business Outcomes")
    add_bullet(document, "Reduction in hidden delay and follow-up dependency.")
    add_bullet(document, "Lower duplicate assignment and ownership confusion.")
    add_bullet(document, "Improved deadline reliability and more honest schedule tracking.")
    add_bullet(document, "Higher visibility into rework sources, quality leakage, and unstable stages.")
    add_bullet(document, "Better departmental coordination during fixture, production, quality, and maintenance handoffs.")
    add_bullet(document, "Stronger confidence for leadership decisions because status is backed by time-stamped operating records.")

    add_heading_band(document, "Future Scalability")
    add_body_paragraph(
        document,
        "PARC is designed to support incremental expansion. New departments can be onboarded through the same control model: department-specific workflows, role-based visibility, assignment templates, escalation rules, report structures, and analytics permissions. This means the platform can grow from a focused operational workflow engine into an enterprise-wide execution governance layer without changing its core operating logic.",
    )
    add_body_paragraph(
        document,
        "Future maturity can include deeper PPC control, stores and dispatch workflows, customer-linked milestone reporting, asset-level analytics, broader machine integration, and richer executive dashboards. The important point for leadership is that the foundation is already correct: accountability, hierarchy, review discipline, and data-backed visibility are built into the operating model from the start.",
    )

    add_heading_band(document, "Conclusion")
    add_body_paragraph(
        document,
        "PARC should be treated as a mission-critical operational control system. Its value is not limited to tracking tasks; its value lies in creating disciplined execution, visible ownership, measurable quality, timely escalation, and trusted management reporting across the organization. When PARC is used consistently, the organization gains control over not only what work exists, but how reliably that work moves from assignment to accepted completion.",
    )
    add_body_paragraph(
        document,
        "Leadership should support PARC because it strengthens governance without disconnecting from actual plant reality. Employees should use PARC correctly because it protects clarity, fairness, and recognition by making work visible and attributable. Departments should rely on PARC because it provides a shared operating language for coordination. In short, PARC is not merely software. It is an operating discipline expressed through software.",
    )


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    build_cover(document)
    build_contents(document)
    populate_sections(document)
    document.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
